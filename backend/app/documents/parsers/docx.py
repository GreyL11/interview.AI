from pathlib import Path

from app.documents.parsers.base import (
    DocumentParser,
    ParseError,
    ensure_text,
    normalize_whitespace,
)
from app.documents.schemas import FileType, NormalizedDocument


class DocxParser(DocumentParser):
    def supports(self, file_type: FileType) -> bool:
        return file_type == FileType.DOCX

    def parse(self, file_path: Path, document_id: str) -> NormalizedDocument:
        import docx

        try:
            document = docx.Document(str(file_path))
        except Exception as exc:
            raise ParseError(f"Could not read DOCX '{file_path.name}': {exc}") from exc

        blocks = [p.text for p in document.paragraphs]

        # Resumes put half their content in tables; dropping them loses the job
        # history entirely.
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))

        text = normalize_whitespace("\n\n".join(b for b in blocks if b.strip()))
        ensure_text(text, file_path)

        heading = next(
            (p.text.strip() for p in document.paragraphs
             if p.style is not None and p.style.name.startswith("Heading") and p.text.strip()),
            "",
        )
        return NormalizedDocument(
            document_id=document_id,
            title=heading or file_path.stem,
            text=text,
            metadata={"parser": "docx", "table_count": len(document.tables)},
        )
