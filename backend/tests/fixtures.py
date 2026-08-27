"""Builders for real document files used by the parser tests.

Files are generated at test time rather than committed as binaries: no opaque
blobs in git, and the fixtures stay readable and adjustable.
"""

from pathlib import Path


def _escape_pdf(text: str) -> str:
    return text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")


def write_pdf(path: Path, lines: list[str]) -> Path:
    """Write a minimal single-page PDF with a text layer.

    Hand-assembled rather than pulling in reportlab — the format is simple
    enough at this scale, and pypdf reads it the same as any other PDF.
    """
    content = ["BT", "/F1 12 Tf", "14 TL", "72 720 Td"]
    for line in lines:
        content.append(f"({_escape_pdf(line)}) Tj")
        content.append("T*")
    content.append("ET")
    stream = "\n".join(content).encode("latin-1")

    objects = [
        b"<</Type/Catalog/Pages 2 0 R>>",
        b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]"
        b"/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>",
        b"<</Length " + str(len(stream)).encode() + b">>\nstream\n" + stream + b"\nendstream",
        b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    ]

    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode() + body + b"\nendobj\n"

    xref_at = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<</Size {len(objects) + 1}/Root 1 0 R>>\n"
        f"startxref\n{xref_at}\n%%EOF\n"
    ).encode()

    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(bytes(out))
    return path


def write_docx(path: Path, paragraphs: list[str], table: list[list[str]] | None = None) -> Path:
    import docx

    document = docx.Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table:
        t = document.add_table(rows=len(table), cols=len(table[0]))
        for row_index, row in enumerate(table):
            for col_index, value in enumerate(row):
                t.cell(row_index, col_index).text = value
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    return path


def write_text(path: Path, text: str, encoding: str = "utf-8") -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding=encoding)
    return path


RESUME_TEXT = """Jane Doe
Senior Data Engineer

Experience

At Acme Corp I built a Kafka ingestion pipeline handling 40k events per second.
I led the migration from batch ETL to streaming, cutting latency from hours to seconds.

Projects

Deduplication service: I implemented a content-hash based dedupe layer in Flink.
"""

TECHNICAL_TEXT = """Database Indexing Notes

A B-tree index speeds up equality and range lookups by keeping keys sorted.
Write amplification is the main cost: every insert must update each index.

Covering indexes let the database answer a query from the index alone.
"""
